"""
Renders a SemanticModel to a Word (.docx) document using python-docx.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from analyzer.quality_checks import QualityFinding
from models.schema import Measure, Relationship, SemanticModel, Table


_GRAY = RGBColor(0xF5, 0xF5, 0xF5)  # light gray for code background
_CODE_FONT = "Courier New"


def render(model: SemanticModel, findings: list[QualityFinding], output_dir: Path) -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    safe_name = model.name.replace(" ", "_").replace("/", "-")
    out_path = output_dir / f"{safe_name}.docx"
    doc = _build(model, findings)
    doc.save(out_path)
    return out_path


def _build(model: SemanticModel, findings: list[QualityFinding]) -> Document:
    doc = Document()
    _set_default_font(doc)

    # Title
    title = doc.add_heading(model.name, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Metadata
    p = doc.add_paragraph()
    p.add_run("Source file: ").bold = True
    p.add_run(model.source_file)
    p = doc.add_paragraph()
    p.add_run("Generated: ").bold = True
    p.add_run(model.generated_at)

    # TOC placeholder
    _add_toc(doc)

    # Overview
    doc.add_heading("Overview", level=1)
    visible_tables = [t for t in model.tables if not t.is_hidden]
    all_measures = [m for t in model.tables for m in t.measures]
    tbl = doc.add_table(rows=5, cols=2)
    tbl.style = "Table Grid"
    _tbl_row(tbl, 0, "Tables", str(len(visible_tables)))
    _tbl_row(tbl, 1, "Relationships", str(len(model.relationships)))
    _tbl_row(tbl, 2, "Total measures", str(len(all_measures)))
    _tbl_row(tbl, 3, "Roles (RLS)", str(len(model.roles)))
    _tbl_row(tbl, 4, "Quality findings", str(len(findings)))
    doc.add_paragraph()

    # Model Diagram
    from renderer.diagram_renderer import build_mermaid
    doc.add_heading("Model Diagram", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "The diagram below is in Mermaid erDiagram format. "
        "It renders automatically in GitHub, VS Code preview, Notion, GitLab, "
        "and most modern Markdown viewers."
    ).italic = True
    _add_code_block(doc, build_mermaid(model))
    doc.add_paragraph()

    # Tables
    doc.add_heading("Tables", level=1)
    for table in model.tables:
        _render_table(doc, table)

    # Relationships
    _render_relationships(doc, model.relationships)

    # Roles
    if model.roles:
        doc.add_heading("Row-Level Security Roles", level=1)
        for role in model.roles:
            doc.add_heading(role.name, level=2)
            if role.table_permissions:
                for perm in role.table_permissions:
                    doc.add_paragraph(perm, style="List Bullet")
            else:
                doc.add_paragraph("No table filters defined.")

    # Quality findings (always show section)
    _render_findings(doc, findings)

    return doc


def _render_table(doc: Document, table: Table) -> None:
    hidden_tag = " (hidden)" if table.is_hidden else ""
    doc.add_heading(f"{table.name}{hidden_tag}", level=2)

    if table.description:
        p = doc.add_paragraph(table.description)
        p.italic = True

    meta = doc.add_paragraph()
    meta.add_run("Role: ").bold = True
    meta.add_run(table.inferred_role + "  ")
    meta.add_run("Type: ").bold = True
    meta.add_run("Calculated table" if table.is_calculated else "Imported/DirectQuery")

    if table.source_query:
        pq = table.source_query
        p = doc.add_paragraph()
        p.add_run("Source: ").bold = True
        p.add_run(pq.source_type)
        if pq.source_details:
            p.add_run(f" — {pq.source_details}")
        p.add_run("  Query complexity: ").bold = True
        p.add_run(pq.complexity_rating)

    # Columns table
    visible_cols = [c for c in table.columns if not c.is_hidden]
    if visible_cols:
        doc.add_heading("Columns", level=3)
        col_tbl = doc.add_table(rows=1 + len(visible_cols), cols=4)
        col_tbl.style = "Table Grid"
        headers = ["Column", "Type", "Calculated", "Description"]
        for i, h in enumerate(headers):
            col_tbl.cell(0, i).text = h
            col_tbl.cell(0, i).paragraphs[0].runs[0].bold = True
        for r, col in enumerate(visible_cols, start=1):
            col_tbl.cell(r, 0).text = col.name
            col_tbl.cell(r, 1).text = col.data_type
            col_tbl.cell(r, 2).text = "Yes" if col.is_calculated else ""
            col_tbl.cell(r, 3).text = col.description or ""
        doc.add_paragraph()

    # Measures
    if table.measures:
        doc.add_heading("Measures", level=3)
        for measure in table.measures:
            _render_measure(doc, measure)

    # Power Query steps
    if table.source_query and table.source_query.step_descriptions:
        pq = table.source_query
        doc.add_heading("Power Query Steps", level=3)
        from extractor.powerquery_parser import _extract_step_names
        step_names = _extract_step_names(pq.m_code)
        for i, desc in enumerate(pq.step_descriptions):
            step_name = step_names[i] if i < len(step_names) else f"Step {i+1}"
            p = doc.add_paragraph(style="List Number")
            p.add_run(step_name + ": ").bold = True
            p.add_run(desc)

        doc.add_heading("M Code", level=4)
        _add_code_block(doc, pq.m_code)

    # Hierarchies
    if table.hierarchies:
        doc.add_heading("Hierarchies", level=3)
        for h in table.hierarchies:
            doc.add_paragraph(f"{h.name}: {' → '.join(h.levels)}", style="List Bullet")


def _render_measure(doc: Document, measure: Measure) -> None:
    hidden_tag = " (hidden)" if measure.is_hidden else ""
    folder_tag = f"  [{measure.display_folder}]" if measure.display_folder else ""
    p = doc.add_heading(f"{measure.name}{hidden_tag}{folder_tag}", level=4)

    if measure.description:
        doc.add_paragraph(measure.description)

    if measure.business_purpose:
        p2 = doc.add_paragraph()
        p2.add_run("Business purpose: ").bold = True
        p2.add_run(measure.business_purpose).italic = True

    meta = doc.add_paragraph()
    if measure.format_string:
        meta.add_run("Format: ").bold = True
        meta.add_run(f"`{measure.format_string}`  ")
    meta.add_run("Complexity: ").bold = True
    meta.add_run(measure.complexity_tier)
    if measure.used_in_visuals:
        meta.add_run(f"  Used in {len(measure.used_in_visuals)} visual(s).")

    _add_code_block(doc, measure.dax)


def _render_relationships(doc: Document, relationships: list[Relationship]) -> None:
    if not relationships:
        return
    doc.add_heading("Relationships", level=1)
    tbl = doc.add_table(rows=1 + len(relationships), cols=5)
    tbl.style = "Table Grid"
    headers = ["From", "To", "Cardinality", "Cross-filter", "Active"]
    for i, h in enumerate(headers):
        tbl.cell(0, i).text = h
        tbl.cell(0, i).paragraphs[0].runs[0].bold = True
    for r, rel in enumerate(relationships, start=1):
        tbl.cell(r, 0).text = f"{rel.from_table}[{rel.from_column}]"
        tbl.cell(r, 1).text = f"{rel.to_table}[{rel.to_column}]"
        tbl.cell(r, 2).text = rel.cardinality.replace("_to_", " \u2192 ").replace("_", " ")
        tbl.cell(r, 3).text = rel.cross_filter_direction
        tbl.cell(r, 4).text = "Yes" if rel.is_active else "No"
    doc.add_paragraph()


def _render_findings(doc: Document, findings: list[QualityFinding]) -> None:
    doc.add_heading("Quality Findings", level=1)
    if not findings:
        doc.add_paragraph("No quality issues found.")
        return
    warnings = [f for f in findings if f.severity == "warning"]
    infos = [f for f in findings if f.severity == "info"]
    for section_findings, label, icon in [
        (warnings, "Warnings", "\u26a0\ufe0f"),
        (infos, "Info", "\u2139\ufe0f"),
    ]:
        if not section_findings:
            continue
        doc.add_heading(f"{label} ({len(section_findings)})", level=2)
        for f in section_findings:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"{icon} ").bold = False
            p.add_run(f.object_name).bold = True
            p.add_run(f" ({f.category}): {f.detail}")


def _add_code_block(doc: Document, code: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = _CODE_FONT
    run.font.size = Pt(9)
    # Light gray shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F5F5F5")
    pPr.append(shd)


def _tbl_row(tbl, row_idx: int, label: str, value: str) -> None:
    tbl.cell(row_idx, 0).text = label
    tbl.cell(row_idx, 0).paragraphs[0].runs[0].bold = True
    tbl.cell(row_idx, 1).text = value


def _add_toc(doc: Document) -> None:
    doc.add_heading("Table of Contents", level=1)
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)
    doc.add_paragraph("(Right-click → Update Field to refresh the table of contents.)")


def _set_default_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
